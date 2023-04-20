import tkinter as tk
import docx
import datetime

def generate_customer_number(name):
    now = datetime.datetime.now()
    customer_number = f"{name[:3].upper()}-{now.strftime('%Y%m%d%H%M%S')}"
    return customer_number

def create_doc():
    # get user inputs from GUI
    name = name_entry.get()
    phone = phone_entry.get()
    email = email_entry.get()
    customer_number = generate_customer_number(name)
    
    # create document and add text
    doc = docx.Document()
    doc.add_heading('Customer Information')
    doc.add_paragraph(f'Name: {name}')
    doc.add_paragraph(f'Phone: {phone}')
    doc.add_paragraph(f'Email: {email}')
    doc.add_paragraph(f'Customer Number: {customer_number}')
    
    # save document
    doc.save(f'{name}.docx')
    
    # clear user inputs from GUI
    name_entry.delete(0, tk.END)
    phone_entry.delete(0, tk.END)
    email_entry.delete(0, tk.END)
    
    # confirm document creation to user
    confirmation_label.config(text='Document created!')

# create GUI window
window = tk.Tk()
window.title('Document Creator')

# create input fields and labels
name_label = tk.Label(window, text='Name:')
name_label.grid(row=0, column=0, padx=10, pady=10)
name_entry = tk.Entry(window)
name_entry.grid(row=0, column=1, padx=10, pady=10)

phone_label = tk.Label(window, text='Phone:')
phone_label.grid(row=1, column=0, padx=10, pady=10)
phone_entry = tk.Entry(window)
phone_entry.grid(row=1, column=1, padx=10, pady=10)

email_label = tk.Label(window, text='Email:')
email_label.grid(row=2, column=0, padx=10, pady=10)
email_entry = tk.Entry(window)
email_entry.grid(row=2, column=1, padx=10, pady=10)

# create button to create document
create_button = tk.Button(window, text='Create Document', command=create_doc)
create_button.grid(row=4, column=0, columnspan=2, padx=10, pady=10)

# create label for confirmation message
confirmation_label = tk.Label(window, text='')
confirmation_label.grid(row=5, column=0, columnspan=2, padx=10, pady=10)

# start GUI
window.mainloop()
